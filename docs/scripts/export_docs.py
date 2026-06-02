"""
export_docs.py — Генерация документации MYCELIUM CORE в DOCX/PDF.

Собирает Markdown-файлы из docs/src/ в единый документ
с оформлением по ГОСТ 7.32-2017 (отчёт о НИР).

Использование:
    python docs/scripts/export_docs.py --lang ru --format docx
    python docs/scripts/export_docs.py --lang en --format pdf
    python docs/scripts/export_docs.py --lang ru --format both

Требования:
    pip install python-docx Pillow
    pip install docx2pdf  # только для PDF, требует MS Word на Windows
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Run: pip install python-docx Pillow")
    sys.exit(1)

# ═══════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════

SCRIPT_DIR = Path(__file__).resolve().parent
DOCS_ROOT = SCRIPT_DIR.parent
SRC_DIR = DOCS_ROOT / "src"
OUTPUT_DIR = DOCS_ROOT / "export"

# ГОСТ 7.32-2017 параметры
GOST_FONT = "Times New Roman"
GOST_FONT_SIZE = Pt(14)
GOST_LINE_SPACING = 1.5
GOST_MARGIN_TOP = Cm(2)
GOST_MARGIN_BOTTOM = Cm(2)
GOST_MARGIN_LEFT = Cm(3)
GOST_MARGIN_RIGHT = Cm(1.5)
GOST_FIRST_LINE_INDENT = Cm(1.25)

# Порядок разделов документации
SECTIONS_RU = [
    ("index.ru.md", "Введение"),
    ("getting-started/installation.ru.md", "Установка"),
    ("getting-started/first-run.ru.md", "Первый запуск"),
    ("getting-started/quick-tour.ru.md", "Быстрый обзор"),
    ("user-guide/overview.ru.md", "Руководство пользователя"),
    ("user-guide/admin-tab.ru.md", "Вкладка Admin"),
    ("user-guide/vote-tab.ru.md", "Вкладка Голосование"),
    ("user-guide/audit-tab.ru.md", "Вкладка Аудит"),
    ("user-guide/logs-tab.ru.md", "Вкладка Логи"),
    ("architecture/overview.ru.md", "Архитектура: обзор"),
    ("architecture/layers.ru.md", "Слои системы"),
    ("architecture/components.ru.md", "Компоненты системы"),
    ("architecture/data-flow.ru.md", "Потоки данных"),
    ("architecture/decisions/adr-001-pyqt6-choice.ru.md", "ADR-001: PyQt6"),
    ("architecture/decisions/adr-002-geth-dev-mode.ru.md", "ADR-002: Geth Dev Mode"),
    ("architecture/decisions/adr-003-ephemeral-chain.ru.md", "ADR-003: Ephemeral Chain"),
    ("architecture/decisions/adr-004-one-session-one-vote.ru.md", "ADR-004: One Session One Vote"),
    ("architecture/decisions/adr-005-i18n-runtime-switch.ru.md", "ADR-005: i18n"),
    ("architecture/decisions/adr-006-layered-architecture.ru.md", "ADR-006: Layered Architecture"),
    ("architecture/decisions/adr-007-error-parser-separation.ru.md", "ADR-007: ErrorParser"),
    ("security/threat-model.ru.md", "Модель угроз (STRIDE)"),
    ("security/sec-checks.ru.md", "Проверки безопасности SEC-01..06"),
    ("security/audit-procedure.ru.md", "Процедура аудита"),
    ("security/known-limitations.ru.md", "Известные ограничения"),
    ("ui-design/design-system.ru.md", "Дизайн-система"),
    ("ui-design/color-palette.ru.md", "Цветовая палитра"),
    ("ui-design/typography.ru.md", "Типографика"),
    ("ui-design/components.ru.md", "Компоненты интерфейса"),
    ("ui-design/icons.ru.md", "Библиотека иконок"),
    ("ui-design/figma-mockups.ru.md", "Макеты Figma"),
    ("deployment/from-source.ru.md", "Сборка из исходников"),
    ("deployment/pyinstaller.ru.md", "Сборка EXE"),
    ("deployment/distribution.ru.md", "Дистрибуция"),
    ("deployment/troubleshooting.ru.md", "Устранение неполадок"),
    ("development/setup.ru.md", "Настройка среды разработки"),
    ("development/testing.ru.md", "Тестирование"),
    ("development/style-guide.ru.md", "Стиль кода"),
    ("development/contributing.ru.md", "Руководство контрибьютора"),
    ("development/git-workflow.ru.md", "Git-процесс"),
    ("reference/glossary.ru.md", "Глоссарий"),
    ("reference/faq.ru.md", "FAQ"),
    ("reference/changelog.ru.md", "Changelog"),
    ("reference/srs.ru.md", "SRS"),
    ("reference/srs-docs.ru.md", "SRS документации"),
    ("reference/license.ru.md", "Лицензия"),
]

SECTIONS_EN = [
    ("index.md", "Introduction"),
    ("getting-started/installation.md", "Installation"),
    ("getting-started/first-run.md", "First Run"),
    ("getting-started/quick-tour.md", "Quick Tour"),
    ("user-guide/overview.md", "User Guide"),
    ("user-guide/admin-tab.md", "Admin Tab"),
    ("user-guide/vote-tab.md", "Vote Tab"),
    ("user-guide/audit-tab.md", "Audit Tab"),
    ("user-guide/logs-tab.md", "Logs Tab"),
    ("architecture/overview.md", "Architecture Overview"),
    ("architecture/layers.md", "System Layers"),
    ("architecture/components.md", "System Components"),
    ("architecture/data-flow.md", "Data Flow"),
    ("architecture/decisions/adr-001-pyqt6-choice.md", "ADR-001: PyQt6"),
    ("architecture/decisions/adr-002-geth-dev-mode.md", "ADR-002: Geth Dev Mode"),
    ("architecture/decisions/adr-003-ephemeral-chain.md", "ADR-003: Ephemeral Chain"),
    ("architecture/decisions/adr-004-one-session-one-vote.md", "ADR-004: One Session One Vote"),
    ("architecture/decisions/adr-005-i18n-runtime-switch.md", "ADR-005: i18n"),
    ("architecture/decisions/adr-006-layered-architecture.md", "ADR-006: Layered Architecture"),
    ("architecture/decisions/adr-007-error-parser-separation.md", "ADR-007: ErrorParser"),
    ("security/threat-model.md", "Threat Model (STRIDE)"),
    ("security/sec-checks.md", "Security Checks SEC-01..06"),
    ("security/audit-procedure.md", "Audit Procedure"),
    ("security/known-limitations.md", "Known Limitations"),
    ("ui-design/design-system.md", "Design System"),
    ("ui-design/color-palette.md", "Color Palette"),
    ("ui-design/typography.md", "Typography"),
    ("ui-design/components.md", "Widget Components"),
    ("ui-design/icons.md", "Icon Library"),
    ("ui-design/figma-mockups.md", "Figma Mockups"),
    ("deployment/from-source.md", "Building from Source"),
    ("deployment/pyinstaller.md", "PyInstaller Build"),
    ("deployment/distribution.md", "Distribution"),
    ("deployment/troubleshooting.md", "Troubleshooting"),
    ("development/setup.md", "Development Setup"),
    ("development/testing.md", "Testing"),
    ("development/style-guide.md", "Style Guide"),
    ("development/contributing.md", "Contributing"),
    ("development/git-workflow.md", "Git Workflow"),
    ("reference/glossary.md", "Glossary"),
    ("reference/faq.md", "FAQ"),
    ("reference/changelog.md", "Changelog"),
    ("reference/srs.md", "SRS"),
    ("reference/srs-docs.md", "Documentation SRS"),
    ("reference/license.md", "License"),
]

# ═══════════════════════════════════════════════════════════
# GOST DOCUMENT SETUP
# ═══════════════════════════════════════════════════════════

def create_gost_document(lang: str) -> Document:
    """Creates a Document with GOST 7.32-2017 formatting."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = GOST_MARGIN_TOP
    section.bottom_margin = GOST_MARGIN_BOTTOM
    section.left_margin = GOST_MARGIN_LEFT
    section.right_margin = GOST_MARGIN_RIGHT

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = GOST_FONT
    font.size = GOST_FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.line_spacing = GOST_LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = GOST_FIRST_LINE_INDENT
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading styles
    for level in range(1, 5):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
        else:
            h_style = doc.styles.add_style(
                style_name, WD_STYLE_TYPE.PARAGRAPH
            )

        h_font = h_style.font
        h_font.name = GOST_FONT
        h_font.bold = True
        h_font.color.rgb = RGBColor(0, 0, 0)

        h_pf = h_style.paragraph_format
        h_pf.line_spacing = GOST_LINE_SPACING
        h_pf.first_line_indent = None
        h_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if level == 1:
            h_font.size = Pt(16)
            h_pf.space_before = Pt(24)
            h_pf.space_after = Pt(12)
            h_pf.page_break_before = True
        elif level == 2:
            h_font.size = Pt(15)
            h_pf.space_before = Pt(18)
            h_pf.space_after = Pt(10)
        elif level == 3:
            h_font.size = Pt(14)
            h_pf.space_before = Pt(14)
            h_pf.space_after = Pt(8)
        else:
            h_font.size = Pt(14)
            h_pf.space_before = Pt(12)
            h_pf.space_after = Pt(6)

    # Code style
    if "Code" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = "Consolas"
        code_font.size = Pt(10)
        code_font.color.rgb = RGBColor(0, 0, 0)
        code_pf = code_style.paragraph_format
        code_pf.line_spacing = 1.0
        code_pf.space_before = Pt(4)
        code_pf.space_after = Pt(4)
        code_pf.first_line_indent = None
        code_pf.left_indent = Cm(1)

    return doc


# ═══════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ═══════════════════════════════════════════════════════════

def parse_markdown_to_docx(doc: Document, md_text: str, section_title: str) -> None:
    """Parses markdown text and adds content to the document."""

    lines = md_text.split("\n")
    in_code_block = False
    in_table = False
    table_rows = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph(line, style="Code")
            i += 1
            continue

        # Empty lines
        if not line.strip():
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            stripped = line.strip()
            # Skip separator rows
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            cells = [
                c.strip()
                for c in stripped.split("|")[1:-1]
            ]
            if cells:
                in_table = True
                table_rows.append(cells)
            i += 1
            continue

        if in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown formatting from headings
            text = _clean_markdown(text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line.strip()):
            i += 1
            continue

        # List items
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.left_indent = Cm(1.25 + (indent // 2) * 0.5)

            marker = list_match.group(2).strip()
            if marker in ("-", "*"):
                prefix = "-- "
            else:
                prefix = f"{marker} "

            _add_formatted_text(p, prefix + text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line.strip())
        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Adds a table to the document."""
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < max_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = _clean_markdown(cell_text)

                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    pf = paragraph.paragraph_format
                    pf.first_line_indent = None
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
                    pf.line_spacing = 1.0

                    for run in paragraph.runs:
                        run.font.size = Pt(11)

    # Bold first row (header)
    if len(rows) > 0:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _add_formatted_text(paragraph, text: str) -> None:
    """Adds text with basic markdown formatting (bold, italic, code)."""
    # Process inline formatting
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(11)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            cleaned = _clean_markdown(part)
            if cleaned:
                paragraph.add_run(cleaned)


def _clean_markdown(text: str) -> str:
    """Removes markdown syntax from text."""
    # Remove links [text](url) -> text
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Remove images ![alt](url)
    text = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"[\1]", text)
    # Remove remaining ** and *
    text = text.replace("**", "").replace("*", "")
    # Remove backticks
    text = text.replace("`", "")
    return text


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

def add_title_page(doc: Document, lang: str) -> None:
    """Adds a GOST-formatted title page."""
    for _ in range(6):
        doc.add_paragraph()

    if lang == "ru":
        title = "MYCELIUM CORE"
        subtitle = (
            "Техническая документация\n"
            "Desktop sandbox-приложение для моделирования,\n"
            "проведения и аудита электронного голосования\n"
            "на локальной приватной Ethereum-сети"
        )
        version_text = "Версия: 1.0.1"
        author_text = "Автор: AlexAgents"
        date_text = f"Дата: {datetime.now().strftime('%d.%m.%Y')}"
    else:
        title = "MYCELIUM CORE"
        subtitle = (
            "Technical Documentation\n"
            "Desktop sandbox application for modeling,\n"
            "executing, and auditing electronic voting\n"
            "on a local private Ethereum network"
        )
        version_text = "Version: 1.0.1"
        author_text = "Author: AlexAgents"
        date_text = f"Date: {datetime.now().strftime('%Y-%m-%d')}"

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = None
    run = p_title.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = GOST_FONT

    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.first_line_indent = None
    run = p_sub.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.name = GOST_FONT

    for _ in range(8):
        doc.add_paragraph()

    for text in (version_text, author_text, date_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = GOST_FONT

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

def add_table_of_contents(doc: Document, lang: str) -> None:
    """Adds a table of contents page."""
    title = "СОДЕРЖАНИЕ" if lang == "ru" else "TABLE OF CONTENTS"
    p = doc.add_heading(title, level=1)
    p.style.paragraph_format.page_break_before = False

    note = (
        "(Обновите оглавление после открытия в MS Word: "
        "ПКМ -> Обновить поле -> Обновить целиком)"
        if lang == "ru"
        else "(Update TOC after opening in MS Word: "
        "Right-click -> Update Field -> Update Entire Table)"
    )
    p_note = doc.add_paragraph(note)
    p_note.paragraph_format.first_line_indent = None
    for run in p_note.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    # Add TOC field
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    fld_char_begin = run._element.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "begin"}
    )
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {})
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = run3._element.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "end"}
    )
    run3._element.append(fld_char_end)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# MAIN EXPORT
# ═══════════════════════════════════════════════════════════

def export_document(lang: str, fmt: str) -> None:
    """Main export function."""
    sections = SECTIONS_RU if lang == "ru" else SECTIONS_EN
    doc = create_gost_document(lang)

    # Title page
    add_title_page(doc, lang)

    # Table of contents
    add_table_of_contents(doc, lang)

    # Sections
    total = len(sections)
    for idx, (file_path, section_title) in enumerate(sections, 1):
        full_path = SRC_DIR / file_path
        if not full_path.exists():
            print(f"  [{idx}/{total}] SKIP (not found): {file_path}")
            continue

        print(f"  [{idx}/{total}] {section_title}: {file_path}")

        md_text = full_path.read_text(encoding="utf-8")

        # Add section heading
        doc.add_heading(section_title, level=1)

        # Parse and add content
        parse_markdown_to_docx(doc, md_text, section_title)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    docx_name = f"MYCELIUM_CORE_docs_{lang}.docx"
    docx_path = OUTPUT_DIR / docx_name
    doc.save(str(docx_path))
    print(f"\n  DOCX saved: {docx_path}")

    if fmt in ("pdf", "both"):
        try:
            from docx2pdf import convert
            pdf_name = f"MYCELIUM_CORE_docs_{lang}.pdf"
            pdf_path = OUTPUT_DIR / pdf_name
            print(f"  Converting to PDF...")
            convert(str(docx_path), str(pdf_path))
            print(f"  PDF saved: {pdf_path}")
        except ImportError:
            print("  WARNING: docx2pdf not installed. Install with:")
            print("    pip install docx2pdf")
            print("  Or open the DOCX in MS Word and save as PDF manually.")
        except Exception as exc:
            print(f"  WARNING: PDF conversion failed: {exc}")
            print("  Open the DOCX in MS Word and save as PDF manually.")

    if fmt == "pdf" and fmt != "both":
        pass  # PDF already attempted above

    print("\nDone!")


# ═══════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="MYCELIUM CORE Documentation Exporter (GOST 7.32-2017)"
    )
    parser.add_argument(
        "--lang",
        choices=["ru", "en"],
        default="ru",
        help="Document language (default: ru)",
    )
    parser.add_argument(
        "--format",
        choices=["docx", "pdf", "both"],
        default="docx",
        help="Output format (default: docx)",
    )
    args = parser.parse_args()

    lang_label = "Russian" if args.lang == "ru" else "English"
    print("=" * 60)
    print(f"  MYCELIUM CORE Documentation Export")
    print(f"  Language: {lang_label}")
    print(f"  Format: {args.format.upper()}")
    print(f"  GOST 7.32-2017 formatting")
    print("=" * 60)
    print()

    export_document(args.lang, args.format)


if __name__ == "__main__":
    main()